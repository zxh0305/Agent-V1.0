#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
工具池模块
==========

本模块实现了智能 Agent 平台的工具池管理功能，包含5个学生高频工具：
1. FileTool - 文件处理工具：PDF文本/表格提取、PDF转Word
2. DataTool - 数据分析工具：Excel/CSV读取、统计计算、图表生成
3. CodeTool - 代码运行工具：安全沙箱运行Python代码
4. PaperTool - 文献查询工具：本地文献库关键词搜索
5. ScheduleTool - 日程管理工具：日程的增删查操作

作者：学生开发团队
版本：1.0.0
"""

import os
import re
import io
import sqlite3
import subprocess
import tempfile
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional, Union
from datetime import datetime
from dataclasses import dataclass

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


# ============================================================
# 基础工具抽象类
# ============================================================

class BaseTool(ABC):
    """
    工具基类

    所有工具必须继承此基类并实现 execute 方法

    Attributes:
        name: 工具名称
        description: 工具描述
        supported_actions: 支持的操作列表
    """

    def __init__(self, name: str, description: str):
        """
        初始化工具

        Args:
            name: 工具名称
            description: 工具描述
        """
        self.name = name
        self.description = description
        self.supported_actions: List[str] = []

    @abstractmethod
    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行工具操作

        Args:
            params: 操作参数字典

        Returns:
            执行结果字典，包含 success 状态和结果数据
        """
        pass

    def validate_params(self, params: Dict[str, Any], required: List[str]) -> bool:
        """
        验证必需参数是否存在

        Args:
            params: 参数字典
            required: 必需参数列表

        Returns:
            验证是否通过
        """
        for key in required:
            if key not in params or params[key] is None:
                return False
        return True


# ============================================================
# 文件处理工具
# ============================================================

class FileTool(BaseTool):
    """
    文件处理工具类

    功能：
    1. PDF 文本提取
    2. PDF 表格提取
    3. PDF 转 Word 文档

    依赖：PyPDF2, python-docx, pdfplumber
    """

    def __init__(self):
        """初始化文件处理工具"""
        super().__init__(
            name="file_tool",
            description="文件处理工具，支持PDF文本/表格提取、PDF转Word转换"
        )
        self.supported_actions = ["extract_text", "extract_tables", "convert_to_word"]

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行文件处理操作

        Args:
            params: 参数字典，包含:
                - action: 操作类型 (extract_text/extract_tables/convert_to_word)
                - file_path: 文件路径
                - output_path: 输出路径（可选）

        Returns:
            执行结果字典
        """
        action = params.get("action", "extract_text")
        file_path = params.get("file_path", "")

        logger.info(f"FileTool 执行操作: {action}, 文件: {file_path}")

        try:
            if action == "extract_text":
                return self._extract_text(file_path)
            elif action == "extract_tables":
                return self._extract_tables(file_path)
            elif action == "convert_to_word":
                output_path = params.get("output_path", "")
                return self._convert_to_word(file_path, output_path)
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }
        except Exception as e:
            logger.error(f"FileTool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        从 PDF 中提取文本

        Args:
            file_path: PDF 文件路径

        Returns:
            提取结果
        """
        try:
            from PyPDF2 import PdfReader

            if not os.path.exists(file_path):
                # 模拟模式：文件不存在时返回模拟数据
                return {
                    "success": True,
                    "message": f"模拟提取 PDF 文本: {file_path}",
                    "data": {
                        "text": "这是模拟提取的PDF文本内容。包含标题、段落和其他文本信息。",
                        "pages": 1,
                        "source": file_path
                    }
                }

            reader = PdfReader(file_path)
            text_content = []
            page_count = len(reader.pages)

            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_content.append({
                        "page": page_num + 1,
                        "text": page_text
                    })

            return {
                "success": True,
                "message": f"成功从 {file_path} 提取文本",
                "data": {
                    "text": "\n".join([p["text"] for p in text_content]),
                    "pages": page_count,
                    "page_details": text_content,
                    "source": file_path
                }
            }

        except ImportError:
            return {
                "success": True,
                "message": f"模拟提取 PDF 文本 (PyPDF2 未安装): {file_path}",
                "data": {
                    "text": "模拟PDF文本内容：本文档包含示例数据和分析结果。",
                    "pages": 1,
                    "source": file_path
                }
            }

    def _extract_tables(self, file_path: str) -> Dict[str, Any]:
        """
        从 PDF 中提取表格

        Args:
            file_path: PDF 文件路径

        Returns:
            提取结果
        """
        try:
            import pdfplumber

            if not os.path.exists(file_path):
                # 模拟模式
                return {
                    "success": True,
                    "message": f"模拟提取 PDF 表格: {file_path}",
                    "data": {
                        "tables": [
                            {
                                "page": 1,
                                "data": [
                                    ["列1", "列2", "列3"],
                                    ["数据1", "100", "200"],
                                    ["数据2", "150", "250"],
                                    ["数据3", "180", "280"]
                                ]
                            }
                        ],
                        "table_count": 1,
                        "source": file_path
                    }
                }

            tables_data = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for table in tables:
                        tables_data.append({
                            "page": page_num + 1,
                            "data": table
                        })

            return {
                "success": True,
                "message": f"成功从 {file_path} 提取 {len(tables_data)} 个表格",
                "data": {
                    "tables": tables_data,
                    "table_count": len(tables_data),
                    "source": file_path
                }
            }

        except ImportError:
            return {
                "success": True,
                "message": f"模拟提取 PDF 表格 (pdfplumber 未安装): {file_path}",
                "data": {
                    "tables": [
                        {
                            "page": 1,
                            "data": [
                                ["姓名", "分数", "等级"],
                                ["张三", "85", "良好"],
                                ["李四", "92", "优秀"]
                            ]
                        }
                    ],
                    "table_count": 1,
                    "source": file_path
                }
            }

    def _convert_to_word(self, file_path: str, output_path: str = "") -> Dict[str, Any]:
        """
        将 PDF 转换为 Word 文档

        Args:
            file_path: PDF 文件路径
            output_path: Word 输出路径

        Returns:
            转换结果
        """
        try:
            from docx import Document

            # 先提取文本
            text_result = self._extract_text(file_path)
            if not text_result["success"]:
                return text_result

            text_content = text_result["data"]["text"]

            # 生成输出路径
            if not output_path:
                base_name = os.path.splitext(os.path.basename(file_path))[0]
                output_path = f"{base_name}_converted.docx"

            # 创建 Word 文档
            doc = Document()
            doc.add_heading("PDF 转换文档", 0)

            # 添加来源信息
            doc.add_paragraph(f"来源文件: {file_path}")
            doc.add_paragraph(f"转换时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph("")

            # 添加内容
            paragraphs = text_content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())

            # 保存文档
            doc.save(output_path)

            return {
                "success": True,
                "message": f"成功将 {file_path} 转换为 Word 文档",
                "data": {
                    "output_path": output_path,
                    "source": file_path,
                    "size": os.path.getsize(output_path) if os.path.exists(output_path) else 0
                }
            }

        except ImportError:
            return {
                "success": True,
                "message": f"模拟 PDF 转 Word (python-docx 未安装): {file_path}",
                "data": {
                    "output_path": output_path or "output.docx",
                    "source": file_path,
                    "size": 0
                }
            }


# ============================================================
# 数据分析工具
# ============================================================

class DataTool(BaseTool):
    """
    数据分析工具类

    功能：
    1. Excel/CSV 数据读取
    2. 统计计算（均值、方差、求和、最大最小值）
    3. 图表生成（柱状图、折线图、饼图、散点图）

    依赖：pandas, matplotlib
    """

    def __init__(self):
        """初始化数据分析工具"""
        super().__init__(
            name="data_tool",
            description="数据分析工具，支持数据读取、统计计算和图表生成"
        )
        self.supported_actions = [
            "read_data", "calculate", "mean", "variance", "sum", "max", "min",
            "generate_chart", "bar", "line", "pie", "scatter"
        ]
        self._current_data = None

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行数据分析操作

        Args:
            params: 参数字典，包含:
                - action/operation: 操作类型
                - file_path: 数据文件路径
                - column: 指定列名
                - chart_type: 图表类型
                - output_path: 图表输出路径

        Returns:
            执行结果字典
        """
        action = params.get("action") or params.get("operation", "read_data")
        file_path = params.get("file_path", "")

        logger.info(f"DataTool 执行操作: {action}, 文件: {file_path}")

        try:
            # 数据读取操作
            if action == "read_data":
                return self._read_data(file_path)

            # 统计计算操作
            elif action in ["mean", "variance", "sum", "max", "min", "calculate"]:
                # 如果有文件路径，先读取数据
                if file_path:
                    read_result = self._read_data(file_path)
                    if not read_result["success"]:
                        return read_result

                column = params.get("column")
                return self._calculate_statistics(action, column, params)

            # 图表生成操作
            elif action in ["generate_chart", "bar", "line", "pie", "scatter"]:
                chart_type = params.get("chart_type", action if action != "generate_chart" else "bar")
                output_path = params.get("output_path", "chart_output.png")

                # 如果有文件路径，先读取数据
                if file_path:
                    read_result = self._read_data(file_path)
                    if not read_result["success"]:
                        return read_result

                return self._generate_chart(chart_type, output_path, params)

            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }

        except Exception as e:
            logger.error(f"DataTool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _read_data(self, file_path: str) -> Dict[str, Any]:
        """
        读取 Excel 或 CSV 数据

        Args:
            file_path: 数据文件路径

        Returns:
            读取结果
        """
        try:
            import pandas as pd

            if not os.path.exists(file_path):
                # 模拟模式：生成模拟数据
                mock_data = pd.DataFrame({
                    "序号": [1, 2, 3, 4, 5],
                    "名称": ["项目A", "项目B", "项目C", "项目D", "项目E"],
                    "数值": [100, 150, 200, 175, 225],
                    "比率": [0.2, 0.3, 0.4, 0.35, 0.45]
                })
                self._current_data = mock_data

                return {
                    "success": True,
                    "message": f"模拟读取数据: {file_path}",
                    "data": {
                        "columns": list(mock_data.columns),
                        "rows": len(mock_data),
                        "preview": mock_data.head().to_dict(),
                        "source": file_path
                    }
                }

            # 根据文件扩展名读取数据
            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            elif file_ext == '.csv':
                df = pd.read_csv(file_path)
            else:
                return {
                    "success": False,
                    "error": f"不支持的文件格式: {file_ext}",
                    "data": None
                }

            self._current_data = df

            return {
                "success": True,
                "message": f"成功读取数据文件: {file_path}",
                "data": {
                    "columns": list(df.columns),
                    "rows": len(df),
                    "preview": df.head().to_dict(),
                    "dtypes": df.dtypes.astype(str).to_dict(),
                    "source": file_path
                }
            }

        except ImportError:
            # pandas 未安装时的模拟数据
            mock_data = {
                "columns": ["序号", "名称", "数值"],
                "rows": 5,
                "preview": {
                    "序号": [1, 2, 3, 4, 5],
                    "名称": ["A", "B", "C", "D", "E"],
                    "数值": [100, 150, 200, 175, 225]
                },
                "source": file_path
            }
            return {
                "success": True,
                "message": f"模拟读取数据 (pandas 未安装): {file_path}",
                "data": mock_data
            }

    def _calculate_statistics(self, operation: str, column: str = None,
                              params: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        计算统计数据

        Args:
            operation: 统计操作类型
            column: 指定列名
            params: 其他参数

        Returns:
            计算结果
        """
        try:
            import pandas as pd
            import numpy as np

            # 获取数据
            if self._current_data is None:
                # 使用模拟数据
                data = pd.DataFrame({
                    "数值": [100, 150, 200, 175, 225]
                })
            else:
                data = self._current_data

            # 确定要计算的列
            numeric_cols = data.select_dtypes(include=[np.number]).columns.tolist()

            if column and column in data.columns:
                calc_data = data[column]
            elif numeric_cols:
                calc_data = data[numeric_cols[0]]
                column = numeric_cols[0]
            else:
                return {
                    "success": False,
                    "error": "没有找到可用于计算的数值列",
                    "data": None
                }

            # 执行计算
            results = {}
            if operation == "mean":
                results["mean"] = float(calc_data.mean())
            elif operation == "variance":
                results["variance"] = float(calc_data.var())
            elif operation == "sum":
                results["sum"] = float(calc_data.sum())
            elif operation == "max":
                results["max"] = float(calc_data.max())
            elif operation == "min":
                results["min"] = float(calc_data.min())
            elif operation == "calculate":
                # 计算所有统计量
                results = {
                    "mean": float(calc_data.mean()),
                    "variance": float(calc_data.var()),
                    "std": float(calc_data.std()),
                    "sum": float(calc_data.sum()),
                    "max": float(calc_data.max()),
                    "min": float(calc_data.min()),
                    "count": int(calc_data.count())
                }

            return {
                "success": True,
                "message": f"成功计算 {column} 列的 {operation} 值",
                "data": {
                    "column": column,
                    "operation": operation,
                    "results": results
                }
            }

        except ImportError:
            # 模拟计算结果
            mock_results = {
                "mean": 170.0,
                "variance": 2062.5,
                "sum": 850.0,
                "max": 225.0,
                "min": 100.0
            }
            return {
                "success": True,
                "message": f"模拟计算 {operation} (pandas/numpy 未安装)",
                "data": {
                    "column": column or "数值",
                    "operation": operation,
                    "results": {operation: mock_results.get(operation, 0)}
                }
            }

    def _generate_chart(self, chart_type: str, output_path: str,
                        params: Dict[str, Any] = None) -> Dict[str, Any]:
        """
        生成图表

        Args:
            chart_type: 图表类型 (bar/line/pie/scatter)
            output_path: 输出文件路径
            params: 其他参数

        Returns:
            生成结果
        """
        try:
            import pandas as pd
            import matplotlib.pyplot as plt
            import matplotlib

            # 设置中文字体
            matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans']
            matplotlib.rcParams['axes.unicode_minus'] = False

            # 获取数据
            if self._current_data is None:
                data = pd.DataFrame({
                    "类别": ["A", "B", "C", "D", "E"],
                    "数值": [100, 150, 200, 175, 225]
                })
            else:
                data = self._current_data

            # 创建图表
            fig, ax = plt.subplots(figsize=(10, 6))

            # 确定 x 和 y 数据
            x_col = params.get("x_column") if params else None
            y_col = params.get("y_column") if params else None

            if x_col and x_col in data.columns:
                x_data = data[x_col]
            else:
                x_data = data.iloc[:, 0]
                x_col = data.columns[0]

            numeric_cols = data.select_dtypes(include=['number']).columns.tolist()
            if y_col and y_col in data.columns:
                y_data = data[y_col]
            elif numeric_cols:
                y_data = data[numeric_cols[0]]
                y_col = numeric_cols[0]
            else:
                y_data = list(range(len(data)))
                y_col = "索引"

            # 根据类型绘制图表
            if chart_type == "bar":
                ax.bar(x_data.astype(str), y_data)
                ax.set_xlabel(x_col)
                ax.set_ylabel(y_col)
                ax.set_title(f"{y_col} 柱状图")

            elif chart_type == "line":
                ax.plot(x_data, y_data, marker='o')
                ax.set_xlabel(x_col)
                ax.set_ylabel(y_col)
                ax.set_title(f"{y_col} 折线图")

            elif chart_type == "pie":
                ax.pie(y_data, labels=x_data.astype(str), autopct='%1.1f%%')
                ax.set_title(f"{y_col} 饼图")

            elif chart_type == "scatter":
                ax.scatter(x_data, y_data)
                ax.set_xlabel(x_col)
                ax.set_ylabel(y_col)
                ax.set_title(f"{x_col} vs {y_col} 散点图")

            plt.tight_layout()
            plt.savefig(output_path, dpi=150, bbox_inches='tight')
            plt.close()

            return {
                "success": True,
                "message": f"成功生成 {chart_type} 图表",
                "data": {
                    "chart_type": chart_type,
                    "output_path": output_path,
                    "x_column": x_col,
                    "y_column": y_col
                }
            }

        except ImportError:
            return {
                "success": True,
                "message": f"模拟生成 {chart_type} 图表 (matplotlib 未安装)",
                "data": {
                    "chart_type": chart_type,
                    "output_path": output_path,
                    "simulated": True
                }
            }


# ============================================================
# 代码运行工具
# ============================================================

class CodeTool(BaseTool):
    """
    代码运行工具类

    功能：
    在沙箱环境中安全运行 Python 代码片段

    安全机制：
    1. 超时限制
    2. 禁止危险模块
    3. 输出长度限制
    """

    # 禁止导入的危险模块
    FORBIDDEN_MODULES = [
        'os.system', 'subprocess', 'shutil.rmtree', 'sys.exit',
        '__import__', 'eval', 'exec', 'compile', 'open',
        'socket', 'requests', 'urllib'
    ]

    # 允许导入的安全模块
    ALLOWED_MODULES = [
        'math', 'random', 'statistics', 'datetime', 'json',
        'collections', 'itertools', 'functools', 're', 'string'
    ]

    def __init__(self, timeout: int = 10, max_output_length: int = 5000):
        """
        初始化代码运行工具

        Args:
            timeout: 代码执行超时时间（秒）
            max_output_length: 最大输出长度
        """
        super().__init__(
            name="code_tool",
            description="代码运行工具，支持安全运行Python代码片段"
        )
        self.supported_actions = ["run", "execute", "validate"]
        self.timeout = timeout
        self.max_output_length = max_output_length

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行代码运行操作

        Args:
            params: 参数字典，包含:
                - code: Python 代码字符串
                - action: 操作类型 (run/validate)

        Returns:
            执行结果字典
        """
        code = params.get("code", "")
        action = params.get("action", "run")

        logger.info(f"CodeTool 执行操作: {action}")

        try:
            if action == "validate":
                return self._validate_code(code)
            elif action in ["run", "execute"]:
                # 先验证代码安全性
                validation = self._validate_code(code)
                if not validation["success"]:
                    return validation
                return self._run_code(code)
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }

        except Exception as e:
            logger.error(f"CodeTool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _validate_code(self, code: str) -> Dict[str, Any]:
        """
        验证代码安全性

        检查代码中是否包含危险操作

        Args:
            code: 代码字符串

        Returns:
            验证结果
        """
        if not code.strip():
            return {
                "success": False,
                "error": "代码不能为空",
                "data": None
            }

        # 检查危险模块和操作
        code_lower = code.lower()
        for forbidden in self.FORBIDDEN_MODULES:
            if forbidden.lower() in code_lower:
                return {
                    "success": False,
                    "error": f"代码包含禁止使用的操作: {forbidden}",
                    "data": {"forbidden_pattern": forbidden}
                }

        # 检查危险语法
        dangerous_patterns = [
            r'__\w+__',  # 双下划线方法
            r'globals\s*\(',
            r'locals\s*\(',
            r'getattr\s*\(',
            r'setattr\s*\(',
            r'delattr\s*\(',
        ]

        for pattern in dangerous_patterns:
            if re.search(pattern, code):
                return {
                    "success": False,
                    "error": f"代码包含潜在危险操作: {pattern}",
                    "data": {"dangerous_pattern": pattern}
                }

        return {
            "success": True,
            "message": "代码验证通过",
            "data": {"validated": True}
        }

    def _run_code(self, code: str) -> Dict[str, Any]:
        """
        在沙箱中运行代码

        Args:
            code: 代码字符串

        Returns:
            运行结果
        """
        # 创建临时文件
        with tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False) as f:
            # 添加安全包装
            wrapped_code = f'''
import sys
from io import StringIO

# 重定向输出
_stdout = sys.stdout
_stderr = sys.stderr
_output = StringIO()
sys.stdout = _output
sys.stderr = _output

try:
{self._indent_code(code)}
except Exception as e:
    print(f"运行错误: {{type(e).__name__}}: {{e}}")
finally:
    sys.stdout = _stdout
    sys.stderr = _stderr
    result = _output.getvalue()
    print(result[:{self.max_output_length}])
'''
            f.write(wrapped_code)
            temp_file = f.name

        try:
            # 使用 subprocess 运行代码
            result = subprocess.run(
                ['python', temp_file],
                capture_output=True,
                text=True,
                timeout=self.timeout
            )

            output = result.stdout + result.stderr
            if len(output) > self.max_output_length:
                output = output[:self.max_output_length] + "\n... (输出已截断)"

            return {
                "success": result.returncode == 0,
                "message": "代码执行完成",
                "data": {
                    "output": output.strip(),
                    "return_code": result.returncode,
                    "execution_time": "< {}s".format(self.timeout)
                }
            }

        except subprocess.TimeoutExpired:
            return {
                "success": False,
                "error": f"代码执行超时 (>{self.timeout}秒)",
                "data": None
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"代码执行失败: {str(e)}",
                "data": None
            }
        finally:
            # 清理临时文件
            if os.path.exists(temp_file):
                os.unlink(temp_file)

    def _indent_code(self, code: str, spaces: int = 4) -> str:
        """
        为代码添加缩进

        Args:
            code: 代码字符串
            spaces: 缩进空格数

        Returns:
            添加缩进后的代码
        """
        indent = ' ' * spaces
        lines = code.split('\n')
        return '\n'.join(indent + line for line in lines)


# ============================================================
# 文献查询工具
# ============================================================

class PaperTool(BaseTool):
    """
    文献查询工具类

    功能：
    本地文献库关键词搜索（模拟实现，避免合规问题）

    注意：本工具为模拟接口，不实际连接知网等付费数据库
    """

    # 模拟文献数据库
    MOCK_PAPERS = [
        {
            "id": "P001",
            "title": "深度学习在自然语言处理中的应用综述",
            "authors": ["张三", "李四"],
            "year": 2023,
            "journal": "计算机学报",
            "keywords": ["深度学习", "自然语言处理", "神经网络", "NLP"],
            "abstract": "本文综述了深度学习技术在自然语言处理领域的最新进展..."
        },
        {
            "id": "P002",
            "title": "基于机器学习的图像识别方法研究",
            "authors": ["王五", "赵六"],
            "year": 2023,
            "journal": "软件学报",
            "keywords": ["机器学习", "图像识别", "卷积神经网络", "CNN"],
            "abstract": "本文提出了一种新的基于机器学习的图像识别方法..."
        },
        {
            "id": "P003",
            "title": "智能Agent系统设计与实现",
            "authors": ["孙七", "周八"],
            "year": 2022,
            "journal": "人工智能学报",
            "keywords": ["智能Agent", "多智能体", "决策系统"],
            "abstract": "本文设计并实现了一个智能Agent系统..."
        },
        {
            "id": "P004",
            "title": "大数据分析技术在教育领域的应用",
            "authors": ["吴九", "郑十"],
            "year": 2023,
            "journal": "教育技术研究",
            "keywords": ["大数据", "教育", "数据分析", "学习分析"],
            "abstract": "本文探讨了大数据分析技术在教育领域的应用前景..."
        },
        {
            "id": "P005",
            "title": "基于Python的数据可视化工具开发",
            "authors": ["陈一", "林二"],
            "year": 2022,
            "journal": "计算机应用研究",
            "keywords": ["Python", "数据可视化", "Matplotlib", "Pandas"],
            "abstract": "本文介绍了使用Python开发数据可视化工具的方法..."
        }
    ]

    def __init__(self):
        """初始化文献查询工具"""
        super().__init__(
            name="paper_tool",
            description="文献查询工具，支持本地文献库关键词搜索"
        )
        self.supported_actions = ["search", "get_detail", "list_all"]

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行文献查询操作

        Args:
            params: 参数字典，包含:
                - action: 操作类型 (search/get_detail/list_all)
                - query: 搜索关键词
                - paper_id: 文献ID

        Returns:
            查询结果字典
        """
        action = params.get("action", "search")
        query = params.get("query", "")

        logger.info(f"PaperTool 执行操作: {action}, 查询: {query}")

        try:
            if action == "search":
                return self._search_papers(query)
            elif action == "get_detail":
                paper_id = params.get("paper_id", "")
                return self._get_paper_detail(paper_id)
            elif action == "list_all":
                return self._list_all_papers()
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }

        except Exception as e:
            logger.error(f"PaperTool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _search_papers(self, query: str) -> Dict[str, Any]:
        """
        搜索文献

        Args:
            query: 搜索关键词

        Returns:
            搜索结果
        """
        if not query.strip():
            return {
                "success": False,
                "error": "搜索关键词不能为空",
                "data": None
            }

        # 关键词匹配搜索
        query_lower = query.lower()
        matched_papers = []

        for paper in self.MOCK_PAPERS:
            score = 0

            # 标题匹配
            if query_lower in paper["title"].lower():
                score += 3

            # 关键词匹配
            for kw in paper["keywords"]:
                if query_lower in kw.lower() or kw.lower() in query_lower:
                    score += 2

            # 摘要匹配
            if query_lower in paper["abstract"].lower():
                score += 1

            if score > 0:
                matched_papers.append({
                    "paper": paper,
                    "relevance_score": score
                })

        # 按相关度排序
        matched_papers.sort(key=lambda x: x["relevance_score"], reverse=True)

        results = [
            {
                "id": p["paper"]["id"],
                "title": p["paper"]["title"],
                "authors": p["paper"]["authors"],
                "year": p["paper"]["year"],
                "journal": p["paper"]["journal"],
                "relevance": p["relevance_score"]
            }
            for p in matched_papers
        ]

        return {
            "success": True,
            "message": f"搜索 '{query}' 完成，找到 {len(results)} 篇相关文献",
            "data": {
                "query": query,
                "total": len(results),
                "papers": results
            }
        }

    def _get_paper_detail(self, paper_id: str) -> Dict[str, Any]:
        """
        获取文献详情

        Args:
            paper_id: 文献ID

        Returns:
            文献详情
        """
        for paper in self.MOCK_PAPERS:
            if paper["id"] == paper_id:
                return {
                    "success": True,
                    "message": f"获取文献 {paper_id} 详情成功",
                    "data": paper
                }

        return {
            "success": False,
            "error": f"未找到文献: {paper_id}",
            "data": None
        }

    def _list_all_papers(self) -> Dict[str, Any]:
        """
        列出所有文献

        Returns:
            文献列表
        """
        papers = [
            {
                "id": p["id"],
                "title": p["title"],
                "authors": p["authors"],
                "year": p["year"]
            }
            for p in self.MOCK_PAPERS
        ]

        return {
            "success": True,
            "message": f"文献库共有 {len(papers)} 篇文献",
            "data": {
                "total": len(papers),
                "papers": papers
            }
        }


# ============================================================
# 日程管理工具
# ============================================================

class ScheduleTool(BaseTool):
    """
    日程管理工具类

    功能：
    1. 添加日程
    2. 查询日程
    3. 删除日程

    存储：SQLite 数据库
    """

    def __init__(self, db_path: str = "schedules.db"):
        """
        初始化日程管理工具

        Args:
            db_path: SQLite 数据库文件路径
        """
        super().__init__(
            name="schedule_tool",
            description="日程管理工具，支持日程的添加、查询和删除"
        )
        self.supported_actions = ["add", "query", "delete", "list", "update"]
        self.db_path = db_path
        self._init_database()

    def _init_database(self) -> None:
        """初始化数据库表"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('''
            CREATE TABLE IF NOT EXISTS schedules (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                content TEXT,
                date TEXT NOT NULL,
                time TEXT,
                reminder INTEGER DEFAULT 0,
                status TEXT DEFAULT 'pending',
                created_at TEXT DEFAULT CURRENT_TIMESTAMP,
                updated_at TEXT DEFAULT CURRENT_TIMESTAMP
            )
        ''')

        conn.commit()
        conn.close()
        logger.info(f"日程数据库初始化完成: {self.db_path}")

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行日程管理操作

        Args:
            params: 参数字典，包含:
                - action: 操作类型 (add/query/delete/list/update)
                - title: 日程标题
                - content: 日程内容
                - date: 日期 (YYYY-MM-DD)
                - time: 时间 (HH:MM)
                - schedule_id: 日程ID

        Returns:
            执行结果字典
        """
        action = params.get("action", "list")

        logger.info(f"ScheduleTool 执行操作: {action}")

        try:
            if action == "add":
                return self._add_schedule(params)
            elif action == "query":
                return self._query_schedules(params)
            elif action == "delete":
                return self._delete_schedule(params)
            elif action == "list":
                return self._list_schedules(params)
            elif action == "update":
                return self._update_schedule(params)
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }

        except Exception as e:
            logger.error(f"ScheduleTool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _add_schedule(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        添加日程

        Args:
            params: 日程参数

        Returns:
            添加结果
        """
        title = params.get("title") or params.get("content", "未命名日程")
        content = params.get("content", "")
        date = params.get("date", datetime.now().strftime("%Y-%m-%d"))
        time = params.get("time", "")
        reminder = params.get("reminder", 0)

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('''
            INSERT INTO schedules (title, content, date, time, reminder)
            VALUES (?, ?, ?, ?, ?)
        ''', (title, content, date, time, reminder))

        schedule_id = cursor.lastrowid
        conn.commit()
        conn.close()

        return {
            "success": True,
            "message": f"日程添加成功",
            "data": {
                "id": schedule_id,
                "title": title,
                "date": date,
                "time": time
            }
        }

    def _query_schedules(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        查询日程

        Args:
            params: 查询参数

        Returns:
            查询结果
        """
        date = params.get("date")
        keyword = params.get("keyword", "")

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        if date:
            cursor.execute('''
                SELECT id, title, content, date, time, status, created_at
                FROM schedules
                WHERE date = ?
                ORDER BY time
            ''', (date,))
        elif keyword:
            cursor.execute('''
                SELECT id, title, content, date, time, status, created_at
                FROM schedules
                WHERE title LIKE ? OR content LIKE ?
                ORDER BY date, time
            ''', (f'%{keyword}%', f'%{keyword}%'))
        else:
            cursor.execute('''
                SELECT id, title, content, date, time, status, created_at
                FROM schedules
                WHERE date >= date('now')
                ORDER BY date, time
                LIMIT 20
            ''')

        rows = cursor.fetchall()
        conn.close()

        schedules = [
            {
                "id": row[0],
                "title": row[1],
                "content": row[2],
                "date": row[3],
                "time": row[4],
                "status": row[5],
                "created_at": row[6]
            }
            for row in rows
        ]

        return {
            "success": True,
            "message": f"查询到 {len(schedules)} 条日程",
            "data": {
                "total": len(schedules),
                "schedules": schedules
            }
        }

    def _delete_schedule(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        删除日程

        Args:
            params: 删除参数

        Returns:
            删除结果
        """
        schedule_id = params.get("schedule_id") or params.get("id")

        if not schedule_id:
            return {
                "success": False,
                "error": "请提供要删除的日程ID",
                "data": None
            }

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('DELETE FROM schedules WHERE id = ?', (schedule_id,))
        deleted = cursor.rowcount > 0

        conn.commit()
        conn.close()

        if deleted:
            return {
                "success": True,
                "message": f"日程 {schedule_id} 删除成功",
                "data": {"deleted_id": schedule_id}
            }
        else:
            return {
                "success": False,
                "error": f"未找到日程: {schedule_id}",
                "data": None
            }

    def _list_schedules(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        列出所有日程

        Args:
            params: 列表参数

        Returns:
            日程列表
        """
        limit = params.get("limit", 50)

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute('''
            SELECT id, title, content, date, time, status, created_at
            FROM schedules
            ORDER BY date DESC, time DESC
            LIMIT ?
        ''', (limit,))

        rows = cursor.fetchall()
        conn.close()

        schedules = [
            {
                "id": row[0],
                "title": row[1],
                "content": row[2],
                "date": row[3],
                "time": row[4],
                "status": row[5],
                "created_at": row[6]
            }
            for row in rows
        ]

        return {
            "success": True,
            "message": f"共有 {len(schedules)} 条日程",
            "data": {
                "total": len(schedules),
                "schedules": schedules
            }
        }

    def _update_schedule(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        更新日程

        Args:
            params: 更新参数

        Returns:
            更新结果
        """
        schedule_id = params.get("schedule_id") or params.get("id")

        if not schedule_id:
            return {
                "success": False,
                "error": "请提供要更新的日程ID",
                "data": None
            }

        # 构建更新语句
        updates = []
        values = []

        if "title" in params:
            updates.append("title = ?")
            values.append(params["title"])
        if "content" in params:
            updates.append("content = ?")
            values.append(params["content"])
        if "date" in params:
            updates.append("date = ?")
            values.append(params["date"])
        if "time" in params:
            updates.append("time = ?")
            values.append(params["time"])
        if "status" in params:
            updates.append("status = ?")
            values.append(params["status"])

        if not updates:
            return {
                "success": False,
                "error": "没有要更新的内容",
                "data": None
            }

        updates.append("updated_at = CURRENT_TIMESTAMP")
        values.append(schedule_id)

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute(f'''
            UPDATE schedules
            SET {", ".join(updates)}
            WHERE id = ?
        ''', values)

        updated = cursor.rowcount > 0
        conn.commit()
        conn.close()

        if updated:
            return {
                "success": True,
                "message": f"日程 {schedule_id} 更新成功",
                "data": {"updated_id": schedule_id}
            }
        else:
            return {
                "success": False,
                "error": f"未找到日程: {schedule_id}",
                "data": None
            }


# ============================================================
# 翻译工具
# ============================================================

class TranslateTool(BaseTool):
    """
    翻译工具类

    功能：
    支持多语言文本翻译（中英互译为主）

    注意：本工具为本地模拟实现，可扩展接入第三方翻译API
    """

    # 常用词汇翻译词典（模拟）
    TRANSLATION_DICT = {
        # 中译英
        "你好": "Hello",
        "谢谢": "Thank you",
        "机器学习": "Machine Learning",
        "深度学习": "Deep Learning",
        "人工智能": "Artificial Intelligence",
        "数据分析": "Data Analysis",
        "自然语言处理": "Natural Language Processing",
        "神经网络": "Neural Network",
        "计算机视觉": "Computer Vision",
        "算法": "Algorithm",
        "模型": "Model",
        "训练": "Training",
        "测试": "Testing",
        "准确率": "Accuracy",
        "损失函数": "Loss Function",
        # 英译中
        "hello": "你好",
        "thank you": "谢谢",
        "machine learning": "机器学习",
        "deep learning": "深度学习",
        "artificial intelligence": "人工智能",
        "data analysis": "数据分析",
        "natural language processing": "自然语言处理",
        "neural network": "神经网络",
        "computer vision": "计算机视觉",
        "algorithm": "算法",
        "model": "模型",
        "training": "训练",
        "testing": "测试",
        "accuracy": "准确率",
        "loss function": "损失函数",
    }

    SUPPORTED_LANGUAGES = ["zh", "en", "ja", "ko", "fr", "de", "es"]

    def __init__(self):
        """初始化翻译工具"""
        super().__init__(
            name="translate_tool",
            description="翻译工具，支持多语言文本翻译（中英互译）"
        )
        self.supported_actions = ["translate", "detect_language", "list_languages"]

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行翻译操作

        Args:
            params: 参数字典，包含:
                - action: 操作类型 (translate/detect_language/list_languages)
                - text: 待翻译文本
                - source_lang: 源语言 (auto/zh/en/...)
                - target_lang: 目标语言 (zh/en/...)

        Returns:
            执行结果字典
        """
        action = params.get("action", "translate")
        text = params.get("text", params.get("query", ""))

        logger.info(f"TranslateTool 执行操作: {action}")

        try:
            if action == "translate":
                source_lang = params.get("source_lang", "auto")
                target_lang = params.get("target_lang", "zh")
                return self._translate(text, source_lang, target_lang)
            elif action == "detect_language":
                return self._detect_language(text)
            elif action == "list_languages":
                return self._list_languages()
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }

        except Exception as e:
            logger.error(f"TranslateTool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _detect_language(self, text: str) -> Dict[str, Any]:
        """
        检测文本语言

        Args:
            text: 待检测文本

        Returns:
            检测结果
        """
        if not text.strip():
            return {
                "success": False,
                "error": "文本不能为空",
                "data": None
            }

        # 简单的语言检测逻辑
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        japanese_chars = len(re.findall(r'[\u3040-\u309f\u30a0-\u30ff]', text))

        total = chinese_chars + english_chars + japanese_chars + 1

        if chinese_chars / total > 0.3:
            detected = "zh"
            lang_name = "中文"
        elif japanese_chars / total > 0.1:
            detected = "ja"
            lang_name = "日文"
        else:
            detected = "en"
            lang_name = "英文"

        return {
            "success": True,
            "message": f"检测到文本语言为: {lang_name}",
            "data": {
                "language": detected,
                "language_name": lang_name,
                "confidence": 0.85
            }
        }

    def _translate(self, text: str, source_lang: str, target_lang: str) -> Dict[str, Any]:
        """
        翻译文本

        Args:
            text: 待翻译文本
            source_lang: 源语言
            target_lang: 目标语言

        Returns:
            翻译结果
        """
        if not text.strip():
            return {
                "success": False,
                "error": "翻译文本不能为空",
                "data": None
            }

        # 自动检测源语言
        if source_lang == "auto":
            detect_result = self._detect_language(text)
            source_lang = detect_result["data"]["language"]

        # 模拟翻译逻辑
        translated_text = text
        text_lower = text.lower()

        # 尝试词典翻译
        for src, tgt in self.TRANSLATION_DICT.items():
            if src.lower() in text_lower:
                translated_text = translated_text.replace(src, tgt)

        # 如果没有匹配到词典，使用简单的模拟翻译
        if translated_text == text:
            if target_lang == "en" and source_lang == "zh":
                # 中译英模拟
                translated_text = f"[Translated to English] {text}"
            elif target_lang == "zh" and source_lang == "en":
                # 英译中模拟
                translated_text = f"[翻译为中文] {text}"
            else:
                translated_text = f"[Translated: {source_lang} -> {target_lang}] {text}"

        return {
            "success": True,
            "message": f"翻译完成 ({source_lang} -> {target_lang})",
            "data": {
                "original_text": text,
                "translated_text": translated_text,
                "source_language": source_lang,
                "target_language": target_lang
            }
        }

    def _list_languages(self) -> Dict[str, Any]:
        """
        列出支持的语言

        Returns:
            语言列表
        """
        languages = {
            "zh": "中文",
            "en": "英文",
            "ja": "日文",
            "ko": "韩文",
            "fr": "法文",
            "de": "德文",
            "es": "西班牙文"
        }

        return {
            "success": True,
            "message": f"支持 {len(languages)} 种语言",
            "data": {
                "languages": languages,
                "total": len(languages)
            }
        }


# ============================================================
# 文本摘要工具
# ============================================================

class SummaryTool(BaseTool):
    """
    文本摘要工具类

    功能：
    1. 文本摘要生成
    2. 关键词提取
    3. 文本统计分析

    使用基于规则和统计的方法实现
    """

    # 停用词列表
    STOP_WORDS = set([
        "的", "了", "和", "是", "在", "我", "有", "这", "个", "们",
        "中", "来", "上", "大", "为", "以", "不", "到", "说", "也",
        "就", "要", "对", "与", "等", "被", "从", "而", "及", "其",
        "the", "a", "an", "is", "are", "was", "were", "be", "been",
        "being", "have", "has", "had", "do", "does", "did", "will",
        "would", "could", "should", "may", "might", "must", "shall",
        "can", "need", "dare", "ought", "used", "to", "of", "in",
        "for", "on", "with", "at", "by", "from", "as", "into", "through"
    ])

    def __init__(self):
        """初始化文本摘要工具"""
        super().__init__(
            name="summary_tool",
            description="文本摘要工具，支持文本摘要、关键词提取和统计分析"
        )
        self.supported_actions = ["summarize", "extract_keywords", "analyze_text"]

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行文本摘要操作

        Args:
            params: 参数字典，包含:
                - action: 操作类型 (summarize/extract_keywords/analyze_text)
                - text: 待处理文本
                - max_length: 摘要最大长度
                - num_keywords: 关键词数量

        Returns:
            执行结果字典
        """
        action = params.get("action", "summarize")
        text = params.get("text", params.get("content", ""))

        logger.info(f"SummaryTool 执行操作: {action}")

        try:
            if action == "summarize":
                max_length = params.get("max_length", 200)
                return self._summarize(text, max_length)
            elif action == "extract_keywords":
                num_keywords = params.get("num_keywords", 5)
                return self._extract_keywords(text, num_keywords)
            elif action == "analyze_text":
                return self._analyze_text(text)
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }

        except Exception as e:
            logger.error(f"SummaryTool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _summarize(self, text: str, max_length: int = 200) -> Dict[str, Any]:
        """
        生成文本摘要

        使用基于句子重要性评分的抽取式摘要方法

        Args:
            text: 原始文本
            max_length: 摘要最大长度

        Returns:
            摘要结果
        """
        if not text.strip():
            return {
                "success": False,
                "error": "文本不能为空",
                "data": None
            }

        # 分句
        sentences = re.split(r'[。！？.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]

        if not sentences:
            return {
                "success": False,
                "error": "无法从文本中提取句子",
                "data": None
            }

        # 计算句子重要性得分
        scored_sentences = []
        for i, sentence in enumerate(sentences):
            score = self._calculate_sentence_score(sentence, i, len(sentences))
            scored_sentences.append((sentence, score))

        # 按得分排序并选择top句子
        scored_sentences.sort(key=lambda x: x[1], reverse=True)

        # 构建摘要
        summary = ""
        selected_sentences = []
        for sentence, score in scored_sentences:
            if len(summary) + len(sentence) <= max_length:
                selected_sentences.append(sentence)
                summary = "。".join(selected_sentences) + "。"
            else:
                break

        if not summary:
            summary = sentences[0][:max_length] + "..."

        return {
            "success": True,
            "message": f"成功生成摘要，原文 {len(text)} 字，摘要 {len(summary)} 字",
            "data": {
                "summary": summary,
                "original_length": len(text),
                "summary_length": len(summary),
                "compression_ratio": round(len(summary) / len(text), 2)
            }
        }

    def _calculate_sentence_score(self, sentence: str, position: int, total: int) -> float:
        """
        计算句子重要性得分

        Args:
            sentence: 句子文本
            position: 句子位置
            total: 总句子数

        Returns:
            重要性得分
        """
        score = 0.0

        # 位置得分：首句和末句权重较高
        if position == 0:
            score += 2.0
        elif position == total - 1:
            score += 1.0
        elif position < total * 0.3:
            score += 0.5

        # 长度得分：适中长度的句子更重要
        length = len(sentence)
        if 20 <= length <= 100:
            score += 1.0
        elif 10 <= length < 20 or 100 < length <= 150:
            score += 0.5

        # 关键词得分
        important_words = ["研究", "发现", "结论", "结果", "表明", "证明",
                          "重要", "关键", "主要", "核心", "首先", "总之"]
        for word in important_words:
            if word in sentence:
                score += 0.5

        return score

    def _extract_keywords(self, text: str, num_keywords: int = 5) -> Dict[str, Any]:
        """
        提取关键词

        使用词频统计和TF思想提取关键词

        Args:
            text: 原始文本
            num_keywords: 关键词数量

        Returns:
            关键词列表
        """
        if not text.strip():
            return {
                "success": False,
                "error": "文本不能为空",
                "data": None
            }

        # 分词（简单实现：按空格和标点分割）
        words = re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z]+', text)

        # 统计词频（排除停用词）
        word_freq = {}
        for word in words:
            word_lower = word.lower()
            if word_lower not in self.STOP_WORDS and len(word) > 1:
                word_freq[word] = word_freq.get(word, 0) + 1

        # 按频率排序
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        keywords = [{"word": w, "frequency": f} for w, f in sorted_words[:num_keywords]]

        return {
            "success": True,
            "message": f"成功提取 {len(keywords)} 个关键词",
            "data": {
                "keywords": keywords,
                "total_words": len(words),
                "unique_words": len(word_freq)
            }
        }

    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """
        分析文本统计信息

        Args:
            text: 原始文本

        Returns:
            统计分析结果
        """
        if not text.strip():
            return {
                "success": False,
                "error": "文本不能为空",
                "data": None
            }

        # 统计信息
        char_count = len(text)
        char_count_no_space = len(text.replace(" ", "").replace("\n", ""))

        # 分词统计
        words = re.findall(r'[\u4e00-\u9fff]+|[a-zA-Z]+', text)
        word_count = len(words)

        # 句子统计
        sentences = re.split(r'[。！？.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        sentence_count = len(sentences)

        # 段落统计
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        paragraph_count = len(paragraphs)

        # 平均句子长度
        avg_sentence_length = char_count_no_space / sentence_count if sentence_count > 0 else 0

        return {
            "success": True,
            "message": "文本分析完成",
            "data": {
                "character_count": char_count,
                "character_count_no_space": char_count_no_space,
                "word_count": word_count,
                "sentence_count": sentence_count,
                "paragraph_count": paragraph_count,
                "average_sentence_length": round(avg_sentence_length, 1)
            }
        }


# ============================================================
# 知识问答工具
# ============================================================

class QATool(BaseTool):
    """
    知识问答工具类

    功能：
    1. 通用知识问答
    2. 学科知识查询
    3. 常识问答

    使用本地知识库实现
    """

    # 知识库（模拟）
    KNOWLEDGE_BASE = {
        "programming": {
            "python": {
                "definition": "Python是一种高级、解释型、通用型编程语言，以简洁清晰的语法著称。",
                "creator": "Guido van Rossum",
                "year": "1991",
                "features": ["简洁易读", "动态类型", "自动内存管理", "丰富的标准库", "跨平台"],
                "use_cases": ["Web开发", "数据分析", "人工智能", "自动化脚本", "科学计算"]
            },
            "java": {
                "definition": "Java是一种广泛使用的面向对象编程语言，具有跨平台特性。",
                "creator": "James Gosling",
                "year": "1995",
                "features": ["跨平台", "面向对象", "安全性高", "多线程", "健壮性"],
                "use_cases": ["企业应用", "Android开发", "大数据", "Web服务", "分布式系统"]
            },
            "javascript": {
                "definition": "JavaScript是一种轻量级的解释型编程语言，是Web开发的核心技术之一。",
                "creator": "Brendan Eich",
                "year": "1995",
                "features": ["动态类型", "事件驱动", "函数式编程", "原型继承"],
                "use_cases": ["前端开发", "后端开发(Node.js)", "移动应用", "游戏开发"]
            }
        },
        "ai": {
            "machine_learning": {
                "definition": "机器学习是人工智能的一个分支，通过算法让计算机从数据中学习规律。",
                "types": ["监督学习", "无监督学习", "强化学习", "半监督学习"],
                "algorithms": ["线性回归", "决策树", "随机森林", "SVM", "神经网络", "K-means"]
            },
            "deep_learning": {
                "definition": "深度学习是机器学习的子集，使用多层神经网络处理复杂模式。",
                "frameworks": ["TensorFlow", "PyTorch", "Keras", "MXNet"],
                "applications": ["图像识别", "语音识别", "自然语言处理", "推荐系统"]
            },
            "nlp": {
                "definition": "自然语言处理(NLP)是人工智能领域，致力于让计算机理解和处理人类语言。",
                "tasks": ["文本分类", "命名实体识别", "情感分析", "机器翻译", "问答系统"]
            }
        },
        "math": {
            "calculus": {
                "definition": "微积分是研究函数的极限、微分、积分以及无穷级数的数学分支。",
                "founders": ["牛顿", "莱布尼茨"],
                "concepts": ["极限", "导数", "积分", "微分方程"]
            },
            "linear_algebra": {
                "definition": "线性代数是研究向量空间和线性映射的数学分支。",
                "concepts": ["矩阵", "向量", "特征值", "线性变换", "行列式"]
            },
            "statistics": {
                "definition": "统计学是收集、分析、解释和展示数据的学科。",
                "concepts": ["均值", "方差", "标准差", "假设检验", "回归分析", "概率分布"]
            }
        }
    }

    # 常见问答对
    FAQ = {
        "什么是人工智能": "人工智能(AI)是计算机科学的一个分支，致力于创建能够模拟人类智能行为的系统，包括学习、推理、感知和决策等能力。",
        "python有什么优点": "Python的主要优点包括：语法简洁易读、学习曲线平缓、拥有丰富的第三方库、跨平台支持、适用于多种应用场景（Web开发、数据科学、AI等）。",
        "如何学习编程": "学习编程的建议：1. 选择一门入门语言（如Python）；2. 学习基础语法和概念；3. 多动手实践编写代码；4. 参与开源项目；5. 阅读优秀代码；6. 持续学习新技术。",
        "什么是数据结构": "数据结构是计算机存储、组织数据的方式，常见的数据结构包括：数组、链表、栈、队列、树、图、哈希表等。",
        "什么是算法": "算法是解决特定问题的一系列明确指令或步骤，具有有限性、确定性、输入、输出和有效性等特征。"
    }

    def __init__(self):
        """初始化知识问答工具"""
        super().__init__(
            name="qa_tool",
            description="知识问答工具，支持编程、AI、数学等领域的知识查询"
        )
        self.supported_actions = ["ask", "search", "list_topics"]

    def execute(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        执行知识问答操作

        Args:
            params: 参数字典，包含:
                - action: 操作类型 (ask/search/list_topics)
                - question: 问题内容
                - topic: 主题分类

        Returns:
            执行结果字典
        """
        action = params.get("action", "ask")
        question = params.get("question", params.get("query", ""))

        logger.info(f"QATool 执行操作: {action}, 问题: {question}")

        try:
            if action == "ask":
                return self._answer_question(question)
            elif action == "search":
                topic = params.get("topic", "")
                return self._search_knowledge(question, topic)
            elif action == "list_topics":
                return self._list_topics()
            else:
                return {
                    "success": False,
                    "error": f"不支持的操作类型: {action}",
                    "data": None
                }

        except Exception as e:
            logger.error(f"QATool 执行出错: {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def _answer_question(self, question: str) -> Dict[str, Any]:
        """
        回答问题

        Args:
            question: 问题内容

        Returns:
            回答结果
        """
        if not question.strip():
            return {
                "success": False,
                "error": "问题不能为空",
                "data": None
            }

        question_lower = question.lower()

        # 先在FAQ中查找
        for q, a in self.FAQ.items():
            if q in question or question in q:
                return {
                    "success": True,
                    "message": "找到匹配的问答",
                    "data": {
                        "question": question,
                        "answer": a,
                        "source": "FAQ",
                        "confidence": 0.9
                    }
                }

        # 在知识库中查找
        answer = self._search_in_knowledge_base(question_lower)
        if answer:
            return {
                "success": True,
                "message": "从知识库中找到答案",
                "data": {
                    "question": question,
                    "answer": answer["content"],
                    "source": answer["source"],
                    "confidence": answer["confidence"]
                }
            }

        # 无法回答
        return {
            "success": True,
            "message": "未能找到确切答案",
            "data": {
                "question": question,
                "answer": f"抱歉，我目前没有关于「{question}」的确切答案。建议您查阅相关专业资料或搜索引擎获取更多信息。",
                "source": "default",
                "confidence": 0.1
            }
        }

    def _search_in_knowledge_base(self, question: str) -> Optional[Dict[str, Any]]:
        """
        在知识库中搜索

        Args:
            question: 问题

        Returns:
            搜索结果
        """
        # 关键词匹配
        keywords = {
            "programming": ["python", "java", "javascript", "编程", "程序", "代码"],
            "ai": ["机器学习", "深度学习", "人工智能", "ai", "ml", "nlp", "自然语言"],
            "math": ["数学", "微积分", "线性代数", "统计", "矩阵", "向量"]
        }

        # 确定领域
        matched_domain = None
        for domain, kws in keywords.items():
            for kw in kws:
                if kw in question:
                    matched_domain = domain
                    break
            if matched_domain:
                break

        if not matched_domain:
            return None

        # 在领域内搜索
        domain_kb = self.KNOWLEDGE_BASE.get(matched_domain, {})
        for topic, info in domain_kb.items():
            topic_keywords = topic.replace("_", " ").split() + [topic]
            for tk in topic_keywords:
                if tk in question:
                    # 构建回答
                    answer_parts = []
                    if "definition" in info:
                        answer_parts.append(info["definition"])
                    if "features" in info:
                        answer_parts.append(f"主要特点：{', '.join(info['features'][:3])}")
                    if "use_cases" in info:
                        answer_parts.append(f"应用场景：{', '.join(info['use_cases'][:3])}")
                    if "concepts" in info:
                        answer_parts.append(f"核心概念：{', '.join(info['concepts'][:4])}")

                    if answer_parts:
                        return {
                            "content": " ".join(answer_parts),
                            "source": f"{matched_domain}/{topic}",
                            "confidence": 0.75
                        }

        return None

    def _search_knowledge(self, query: str, topic: str = "") -> Dict[str, Any]:
        """
        搜索知识库

        Args:
            query: 搜索查询
            topic: 主题限定

        Returns:
            搜索结果
        """
        results = []

        search_domains = [topic] if topic and topic in self.KNOWLEDGE_BASE else self.KNOWLEDGE_BASE.keys()

        for domain in search_domains:
            domain_kb = self.KNOWLEDGE_BASE.get(domain, {})
            for topic_name, info in domain_kb.items():
                # 检查是否匹配
                topic_text = f"{topic_name} {info.get('definition', '')}"
                if query.lower() in topic_text.lower():
                    results.append({
                        "domain": domain,
                        "topic": topic_name,
                        "definition": info.get("definition", ""),
                        "relevance": 0.8
                    })

        return {
            "success": True,
            "message": f"搜索 '{query}' 找到 {len(results)} 条结果",
            "data": {
                "query": query,
                "total": len(results),
                "results": results
            }
        }

    def _list_topics(self) -> Dict[str, Any]:
        """
        列出所有知识主题

        Returns:
            主题列表
        """
        topics = {}
        for domain, domain_kb in self.KNOWLEDGE_BASE.items():
            topics[domain] = list(domain_kb.keys())

        return {
            "success": True,
            "message": "知识库主题列表",
            "data": {
                "topics": topics,
                "domains": list(self.KNOWLEDGE_BASE.keys())
            }
        }


# ============================================================
# 工具管理器
# ============================================================

class ToolManager:
    """
    工具管理器类

    统一管理和调用所有工具

    Attributes:
        tools: 工具实例字典
        enabled_tools: 启用的工具列表
    """

    def __init__(self, db_path: str = "schedules.db"):
        """
        初始化工具管理器

        Args:
            db_path: 日程数据库路径
        """
        self.tools: Dict[str, BaseTool] = {}
        self.enabled_tools: List[str] = []
        self._init_tools(db_path)

    def _init_tools(self, db_path: str) -> None:
        """
        初始化所有工具

        Args:
            db_path: 日程数据库路径
        """
        # 注册所有工具（原有5个 + 新增3个）
        self.register_tool(FileTool())
        self.register_tool(DataTool())
        self.register_tool(CodeTool())
        self.register_tool(PaperTool())
        self.register_tool(ScheduleTool(db_path))
        # 新增工具
        self.register_tool(TranslateTool())
        self.register_tool(SummaryTool())
        self.register_tool(QATool())

        # 默认启用所有工具
        self.enabled_tools = list(self.tools.keys())

        logger.info(f"工具管理器初始化完成，已注册 {len(self.tools)} 个工具")

    def register_tool(self, tool: BaseTool) -> None:
        """
        注册工具

        Args:
            tool: 工具实例
        """
        self.tools[tool.name] = tool
        logger.info(f"注册工具: {tool.name}")

    def unregister_tool(self, tool_name: str) -> bool:
        """
        注销工具

        Args:
            tool_name: 工具名称

        Returns:
            是否注销成功
        """
        if tool_name in self.tools:
            del self.tools[tool_name]
            if tool_name in self.enabled_tools:
                self.enabled_tools.remove(tool_name)
            logger.info(f"注销工具: {tool_name}")
            return True
        return False

    def enable_tool(self, tool_name: str) -> bool:
        """
        启用工具

        Args:
            tool_name: 工具名称

        Returns:
            是否启用成功
        """
        if tool_name in self.tools and tool_name not in self.enabled_tools:
            self.enabled_tools.append(tool_name)
            return True
        return False

    def disable_tool(self, tool_name: str) -> bool:
        """
        禁用工具

        Args:
            tool_name: 工具名称

        Returns:
            是否禁用成功
        """
        if tool_name in self.enabled_tools:
            self.enabled_tools.remove(tool_name)
            return True
        return False

    def set_enabled_tools(self, tool_names: List[str]) -> None:
        """
        设置启用的工具列表

        Args:
            tool_names: 工具名称列表
        """
        self.enabled_tools = [name for name in tool_names if name in self.tools]

    def run_tool(self, tool_name: str, params: Dict[str, Any]) -> Dict[str, Any]:
        """
        运行指定工具

        Args:
            tool_name: 工具名称
            params: 工具参数

        Returns:
            执行结果字典
        """
        # 检查工具是否存在
        if tool_name not in self.tools:
            return {
                "success": False,
                "error": f"工具不存在: {tool_name}",
                "data": None
            }

        # 检查工具是否启用
        if tool_name not in self.enabled_tools:
            return {
                "success": False,
                "error": f"工具未启用: {tool_name}",
                "data": None
            }

        # 执行工具
        tool = self.tools[tool_name]
        logger.info(f"执行工具: {tool_name}, 参数: {params}")

        try:
            result = tool.execute(params)
            return result
        except Exception as e:
            logger.error(f"工具执行出错: {tool_name} - {str(e)}")
            return {
                "success": False,
                "error": str(e),
                "data": None
            }

    def get_tool_info(self, tool_name: str = None) -> Dict[str, Any]:
        """
        获取工具信息

        Args:
            tool_name: 工具名称，为空则返回所有工具信息

        Returns:
            工具信息字典
        """
        if tool_name:
            if tool_name in self.tools:
                tool = self.tools[tool_name]
                return {
                    "name": tool.name,
                    "description": tool.description,
                    "supported_actions": tool.supported_actions,
                    "enabled": tool_name in self.enabled_tools
                }
            return None

        # 返回所有工具信息
        return {
            name: {
                "description": tool.description,
                "supported_actions": tool.supported_actions,
                "enabled": name in self.enabled_tools
            }
            for name, tool in self.tools.items()
        }

    def list_tools(self) -> List[Dict[str, Any]]:
        """
        列出所有工具

        Returns:
            工具列表
        """
        return [
            {
                "name": tool.name,
                "description": tool.description,
                "enabled": tool.name in self.enabled_tools
            }
            for tool in self.tools.values()
        ]


# 便捷函数：创建工具管理器实例
def create_tool_manager(db_path: str = "schedules.db") -> ToolManager:
    """
    创建工具管理器实例

    Args:
        db_path: 日程数据库路径

    Returns:
        ToolManager 实例
    """
    return ToolManager(db_path)


# 测试代码
if __name__ == "__main__":
    # 创建工具管理器
    tm = create_tool_manager()

    print("=" * 60)
    print("工具池测试")
    print("=" * 60)

    # 列出所有工具
    print("\n已注册的工具:")
    for tool_info in tm.list_tools():
        status = "✓" if tool_info["enabled"] else "✗"
        print(f"  {status} {tool_info['name']}: {tool_info['description']}")

    # 测试各工具
    print("\n" + "-" * 40)
    print("测试文件处理工具:")
    result = tm.run_tool("file_tool", {"action": "extract_text", "file_path": "test.pdf"})
    print(f"  结果: {result['message']}")

    print("\n" + "-" * 40)
    print("测试数据分析工具:")
    result = tm.run_tool("data_tool", {"action": "mean", "file_path": "data.csv"})
    print(f"  结果: {result['message']}")

    print("\n" + "-" * 40)
    print("测试代码运行工具:")
    result = tm.run_tool("code_tool", {"action": "run", "code": "print('Hello, Agent!')"})
    print(f"  结果: {result['message']}")

    print("\n" + "-" * 40)
    print("测试文献查询工具:")
    result = tm.run_tool("paper_tool", {"action": "search", "query": "机器学习"})
    print(f"  结果: {result['message']}")

    print("\n" + "-" * 40)
    print("测试日程管理工具:")
    result = tm.run_tool("schedule_tool", {
        "action": "add",
        "title": "测试日程",
        "date": "2024-01-15",
        "time": "10:00"
    })
    print(f"  结果: {result['message']}")
